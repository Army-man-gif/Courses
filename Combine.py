import tkinter as tk
from tkinter import filedialog
import os
from PyPDF2 import PdfReader, PdfWriter
from docx import Document

root = tk.Tk()
root.withdraw()


file_path1 = filedialog.askopenfilename(
    title="Select first file",
    filetypes=(("Text files", "*.txt"),
               ("Word documents", "*.docx"),
               ("PDF files", "*.pdf"))
)

file_path2 = filedialog.askopenfilename(
    title="Select second file",
    filetypes=(("Text files", "*.txt"),
               ("Word documents", "*.docx"),
               ("PDF files", "*.pdf"))
)

fileType = ""
if not (file_path1 and file_path2):
    print("Both files must be selected")
    exit()
directory1 = os.path.dirname(file_path1)
directory2 = os.path.dirname(file_path2)

_, ext1 = os.path.splitext(file_path1)
_, ext2 = os.path.splitext(file_path2)
ext1 = ext1.lower()
ext2 = ext2.lower()

if ext1 != ext2:
    print("Both files must be the same type")
    exit()
    
if ext1 == ".txt":
    fileType = ".txt"
elif ext1 == ".docx":
    fileType = ".docx"
elif ext1 == ".pdf":
    fileType=".pdf"
else:
    print("Unknown file type:", ext1)
        
if(fileType == ".docx"):
    document1 = Document(file_path1)
    document2 = Document(file_path2)       
    combined_doc = Document()
    for doc in [document1, document2]:
        for element in doc.element.body:
            combined_doc.element.body.append(element)
    combined_doc.save(os.path.join(directory1,"combined.docx"))
        
elif(fileType == ".pdf"):       
    pdf1 = PdfReader(file_path1)
    pdf2 = PdfReader(file_path2)
    writer = PdfWriter()
    for page in pdf1.pages + pdf2.pages:
        writer.add_page(page)
    with open(os.path.join(directory1,'combined.pdf'), "wb") as f:
        writer.write(f)
elif(fileType == ".txt"):
    with open(file_path1, "r") as f1, open(file_path2, "r") as f2:
        content1 = f1.read()
        content2 = f2.read()
        finalFilePath = os.path.join(directory1,'combined.txt')
        with open(finalFilePath, 'w') as destinationFile:
            destinationFile.write(content1 + content2)
else:
    print("Something wrong with one of your file selections")