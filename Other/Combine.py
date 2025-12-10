import tkinter as tk
from tkinter import filedialog
import os
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
from docxcompose.composer import Composer

root = tk.Tk()
root.withdraw()


file_path1 = filedialog.askopenfilename(
    title="Select first file",
    filetypes=(("Word documents", "*.docx"),
               ("Text files", "*.txt"),
               ("PDF files", "*.pdf"))
)

file_path2 = filedialog.askopenfilename(
    title="Select second file",
    filetypes=(("Word documents", "*.docx"),
               ("Text files", "*.txt"),
               ("PDF files", "*.pdf"))
)

fileType = ""
if not (file_path1 and file_path2):
    print("Both files must be selected")
    exit()
directory1 = os.path.dirname(file_path1)
directory2 = os.path.dirname(file_path2)

_, ext1 = os.path.splitext(file_path1)
_, ext2 = os.path.splitext(file_path2)
ext1 = ext1.lower()
ext2 = ext2.lower()

if ext1 != ext2:
    print("Both files must be the same type")
    exit()
    
if ext1 == ".txt":
    fileType = ".txt"
elif ext1 == ".docx":
    fileType = ".docx"
elif ext1 == ".pdf":
    fileType=".pdf"
else:
    print("Unknown file type:", ext1)
 
'''
if(fileType == ".docx"):

    def copy_headers_footers(source_doc, target_doc):
        try:
            source_section = source_doc.sections[0]  # first section usually has headers/footers
            for section in target_doc.sections:
                # Break link to previous to allow independent headers/footers
                section.header.is_linked_to_previous = False
                section.footer.is_linked_to_previous = False
                # Copy header paragraphs
                for paragraph in source_section.header.paragraphs:
                    section.header.add_paragraph(paragraph.text)
                # Copy footer paragraphs
                for paragraph in source_section.footer.paragraphs:
                    section.footer.add_paragraph(paragraph.text)
        except IndexError:
            # No headers/footers in source document
            pass

    def append_paragraphs_with_style(source_doc, target_doc):
        for paragraph in source_doc.paragraphs:
            # Add paragraph with same style
            new_para = target_doc.add_paragraph(paragraph.text, style=paragraph.style)
            # Copy runs for bold, italics, underline, etc.
            for i, run in enumerate(paragraph.runs):
                if i < len(new_para.runs):
                    new_run = new_para.runs[i]
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.name = run.font.name
                    new_run.font.size = run.font.size
    document1 = Document(file_path1)
    document2 = Document(file_path2)       
    combined_doc = Document()
    for doc in [document1, document2]:
        append_paragraphs_with_style(doc, combined_doc)

    copy_headers_footers(document1, combined_doc)
    combined_doc.save(os.path.join(directory1,"combined.docx"))
''' 
'''
if(fileType == ".docx"):
    document1 = Document(file_path1)
    document2 = Document(file_path2)       
    combined_doc = Document()
    for doc in [document1, document2]:
        for element in doc.element.body:
            combined_doc.element.body.append(element)

    # Copy headers and footers from document1 if it has them
    try:
        source_section = document1.sections[0]
        for section in combined_doc.sections:
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
            # Copy header
            for paragraph in source_section.header.paragraphs:
                section.header.add_paragraph(paragraph.text)
            # Copy footer
            for paragraph in source_section.footer.paragraphs:
                section.footer.add_paragraph(paragraph.text)
    except IndexError:
        pass  # no headers/footers in document1        
    combined_doc.save(os.path.join(directory1,"combined.docx"))
'''

if(fileType == ".docx"):
    doc1 = Document(file_path1)
    doc2 = Document(file_path2)

    composer = Composer(doc1)
    composer.append(doc2)
    composer.save(os.path.join(directory1,"combined.docx"))
elif(fileType == ".pdf"):       
    pdf1 = PdfReader(file_path1)
    pdf2 = PdfReader(file_path2)
    writer = PdfWriter()
    for page in pdf1.pages:
        writer.add_page(page)
    for page in pdf2.pages:
        writer.add_page(page)
    with open(os.path.join(directory1,'AKhaitan-2025_Oct_12_AK_CV_with_cover_letter.pdf'), "wb") as f:
        writer.write(f)
elif(fileType == ".txt"):
    with open(file_path1, "r") as f1, open(file_path2, "r") as f2:
        content1 = f1.read()
        content2 = f2.read()
        finalFilePath = os.path.join(directory1,'combined.txt')
        with open(finalFilePath, 'w') as destinationFile:
            destinationFile.write(content1 + content2)
else:
    print("Something wrong with one of your file selections")